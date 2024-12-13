import sys
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QDateEdit, QGridLayout, QMessageBox, QFileDialog
)
from PyQt5.QtCore import QDate
from docx import Document
from datetime import datetime
from dateutil.relativedelta import relativedelta
import sys
from PyQt5.QtCore import QDate, QThread, pyqtSignal
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QGridLayout, QLabel, QLineEdit, QDateEdit,
    QHBoxLayout, QPushButton, QFileDialog, QMessageBox, QTextEdit
)
import os
import requests
from docx.shared import Pt


BITRIX_WEBHOOK_URL = "https://pravburo.bitrix24.ru/rest/33/16jfafz5traurhir/"
HEADERS = {"Content-Type": "application/json"}

#--------------------------------------------------------------------------------------------------------------------------------------------------------------
def number_to_words(num):
    if num == 0:
        return "ноль"

    units = (
        "", "один", "два", "три", "четыре", "пять",
        "шесть", "семь", "восемь", "девять"
    )
    teens = (
        "десять", "одиннадцать", "двенадцать", "тринадцать",
        "четырнадцать", "пятнадцать", "шестнадцать",
        "семнадцать", "восемнадцать", "девятнадцать"
    )
    tens = (
        "", "", "двадцать", "тридцать", "сорок",
        "пятьдесят", "шестьдесят", "семьдесят",
        "восемьдесят", "девяносто"
    )
    hundreds = (
        "", "сто", "двести", "триста", "четыреста",
        "пятьсот", "шестьсот", "семьсот", "восемьсот",
        "девятьсот"
    )
    thousands_forms = ("тысяча", "тысячи", "тысяч")
    millions_forms = ("миллион", "миллиона", "миллионов")

    def get_form(number, forms):
        if 11 <= number % 100 <= 19:
            return forms[2]
        elif number % 10 == 1:
            return forms[0]
        elif 2 <= number % 10 <= 4:
            return forms[1]
        else:
            return forms[2]

    def three_digit_number_to_words(n):
        result = []
        if n >= 100:
            result.append(hundreds[n // 100])
            n %= 100
        if 10 <= n < 20:
            result.append(teens[n - 10])
        else:
            if n >= 20:
                result.append(tens[n // 10])
            if n % 10 > 0:
                result.append(units[n % 10])
        return " ".join(result).strip()

    result = []

    if num >= 1_000_000:
        millions = num // 1_000_000
        result.append(f"{three_digit_number_to_words(millions)} {get_form(millions, millions_forms)}")
        num %= 1_000_000

    if num >= 1_000:
        thousands = num // 1_000
        thousands_text = three_digit_number_to_words(thousands)
        # Меняем "один" и "два" на женский род в тысячах
        thousands_text = thousands_text.replace("один", "одна").replace("два", "две")
        result.append(f"{thousands_text} {get_form(thousands, thousands_forms)}")
        num %= 1_000

    if num > 0:
        result.append(three_digit_number_to_words(num))

    return " ".join(result).strip()
#----------------------------------------------------------------------------------------------------------------------------------------------------------

def resource_path(relative_path):
    """Получает абсолютный путь к ресурсам, поддерживает PyInstaller"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)



def get_second_payment(num_payments, total_amount, discount, start_date):
    payments_table = calculate_payments(num_payments, total_amount, discount, start_date)
    
    if len(payments_table) > 1:
        return payments_table[1]  
    else:
        return None  


def fill_template_text_only(template_path, output_path, data):
    """
    Заполняет текстовые параграфы и таблицы в шаблоне Word документа данными.
    Устанавливает размер текста 10 пт.
    Выделяет жирным второе вхождение ключа "ФИО".
    """
    doc = Document(template_path)
    font_size = Pt(10)
    count_fio = 0  # Счётчик для отслеживания количества вхождений "ФИО"

    # Обработка параграфов
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
        
        for run in paragraph.runs:
            # Обработка ключа "инициалы"
            if "инициалы" in data and data["инициалы"].lower() in run.text.lower():
                run.font.bold = True
            
            # Обработка второго вхождения "ФИО"
            if "ФИО" in data and data["ФИО"].lower() in run.text:
                count_fio += 1
                if count_fio == 2:  # Второе вхождение
                    run.font.bold = True
            
            run.font.size = font_size

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"  
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                    
                    for run in paragraph.runs:
                        # Обработка ключа "инициалы"
                        if "инициалы" in data and data["инициалы"].lower() in run.text.lower():
                            run.font.bold = True
                        
                        if "ФИО" in data and data["ФИО"].lower() in run.text.lower():
                            run.font.bold = True
                        
                        run.font.size = font_size

    # Удаление строк с "0 рублей"
    for table in doc.tables:
        for row in list(table.rows): 
            first_cell_text = row.cells[0].text.strip()  
            if first_cell_text == "0 рублей":
                table._element.remove(row._element)

    doc.save(output_path)


    


def find_deal_by_title(title):
    """
    Ищет сделку по полю TITLE.
    :param title: Название сделки (соответствует "ФИО").
    :return: ID сделки или None, если сделка не найдена.
    """
    url = f"{BITRIX_WEBHOOK_URL}crm.deal.list"
    params = {
        "filter": {"TITLE": title},  
        "select": ["ID"],  
    }
    response = requests.post(url, json=params, headers=HEADERS)
    response_data = response.json()

    if response.status_code == 200 and "result" in response_data and response_data["result"]:
        return response_data["result"][0]["ID"]
    return None
 

def insert_table_after_heading(doc_path, rows_count, table_data):
    """
    Вставляет таблицу после заголовка с пропущенной строкой в существующий документ.
    Устанавливает размер текста 10 пт.
    
    :param doc_path: Путь к существующему документу, в который нужно добавить таблицу.
    :param rows_count: Количество строк в таблице (не включая заголовки).
    :param table_data: Данные для заполнения таблицы.
    """
    heading_text = "ГРАФИК ПЛАТЕЖЕЙ"
    
    doc = Document(doc_path)
    
    font_size = Pt(10)  

    for paragraph in doc.paragraphs:
        if heading_text in paragraph.text:
            doc.add_paragraph()

            table = doc.add_table(rows=rows_count + 1, cols=3)
            #table.style = 'Table Grid' ХУЕТА
            
            headers = ["П/П", "Дата платежа", "Сумма платежа"]
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = font_size

            for row_idx, row_data in enumerate(table_data, start=1):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(cell_data)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = font_size
            table.style = 'Table Grid'
            paragraph._element.addnext(table._element)
            break
            
    else:
        print(f"Заголовок '{heading_text}' не найден.")
    
    doc.save(doc_path)
    print(f"Документ успешно сохранён: {doc_path}")
    
    
    

def calculate_payments(num_payments, total_amount, discount, start_date, first_payment):
    # Вычисляем оставшуюся сумму для расчета регулярных платежей
    remaining_amount = (total_amount - discount) - first_payment
    remaining_payments = num_payments - 1  # Уменьшаем количество платежей на 1

    if remaining_payments <= 0:
        raise ValueError("Количество платежей должно быть больше 1 при наличии первого платежа.")

    # Рассчитываем регулярный платеж
    payment_amount = remaining_amount / remaining_payments
    payment_amount_rounded = round(payment_amount, -2)
    
    # Корректируем последний платеж на разницу округления
    total_rounded = payment_amount_rounded * remaining_payments
    difference = remaining_amount - total_rounded

    # Преобразуем дату начала платежей
    if isinstance(start_date, str):
        start_date = datetime.strptime(start_date, "%d.%m.%Y")

    table_data = []

    # Добавляем первый платеж с сегодняшней датой
    today = datetime.today()
    table_data.append([1, today.strftime("%d.%m.%Y"), f"{first_payment:.2f}"])

    # Генерируем график для оставшихся платежей
    current_date = start_date
    for i in range(remaining_payments):
        current_date += relativedelta(months=1)

        # Учет февраля
        if current_date.month == 2 and current_date.day > 28:
            current_date = current_date.replace(day=28)
        else:
            try:
                current_date = current_date.replace(day=start_date.day)
            except ValueError:
                current_date = current_date.replace(day=1) + relativedelta(day=31)

        # Учет разницы для последнего платежа
        payment = payment_amount_rounded
        if i == remaining_payments - 1:
            payment += difference

        table_data.append([i + 2, current_date.strftime("%d.%m.%Y"), f"{payment:.2f}"])

    return table_data


class WorkerThread(QThread):
    log_signal = pyqtSignal(str)
    success_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, data, template_path, output_path, discount, num_payments, start_date):
        super().__init__()
        self.data = data
        self.template_path = template_path
        self.output_path = output_path
        self.discount = discount
        self.num_payments = num_payments
        self.start_date = start_date

    def run(self):
        try:
            today = datetime.now()

            formatted_date = today.strftime("«%d» %B %Y г.")

            months = {
                "January": "января", "February": "февраля", "March": "марта",
                "April": "апреля", "May": "мая", "June": "июня",
                "July": "июля", "August": "августа", "September": "сентября",
                "October": "октября", "November": "ноября", "December": "декабря"
            }
            month_name = today.strftime("%B")
            formatted_date = formatted_date.replace(month_name, months[month_name])

            self.data["today"] = formatted_date
            
            self.log_signal.emit("Начало генерации документа.")
            if self.data['сумма бонус'] == '':
                self.data['сумма бонус'] = '0'
            fio_parts = self.data["ФИО"].split()
            if len(fio_parts) != 3:
                raise ValueError("Поле 'ФИО' должно содержать ровно три слова.")
            self.data["инициалы"] = f"{fio_parts[0]} {fio_parts[1][0]}. {fio_parts[2][0]}."
            self.data['сумма юристы'] = str(int(self.data['сумма юристы']) - int(self.discount))
            self.data["сумма"] = str(int(self.data["сумма бонус"]) + int(self.data["сумма юристы"]))
            self.data["words_sum"] = number_to_words(int(self.data["сумма юристы"]))

            self.log_signal.emit("Заполнение шаблона...")
            fill_template_text_only(self.template_path, self.output_path, self.data)

            self.log_signal.emit("Вставка таблицы...")
            total = int(self.data["сумма"]) + self.discount
            table_data = calculate_payments(self.num_payments, total, self.discount, self.start_date, int(self.data['Первый платеж']))
            print(table_data)
            insert_table_after_heading(self.output_path, self.num_payments, table_data)
            
            deal_id = find_deal_by_title(self.data['ФИО'])
            if deal_id:
                payment = get_second_payment(total_amount=int(self.data['сумма бонус']) + int(self.data['сумма юристы']),num_payments=int(self.num_payments),discount=int(self.discount), start_date=datetime.today())[2]
                
                try:
                    url = f"{BITRIX_WEBHOOK_URL}crm.deal.update"
                    params = {
                                "id": deal_id,
                                "fields": {
                                    'UF_CRM_1712048804382': self.data["дата выдачи"],  
                                    'UF_CRM_1712048830962': self.data["кем"],          
                                    'UF_CRM_1712048849415': self.data["код"],         
                                    'UF_CRM_1712056472656': self.data["место рождения"],  
                                    'UF_CRM_1732785003047': f"{int(total)}|RUB",  
                                    'UF_CRM_1732785067451': f"{int(self.discount)}|RUB",  
                                    'UF_CRM_1732785118555': datetime.today().strftime("%Y-%m-%d %H:%M:%S"), 
                                    'UF_CRM_1732785152659': f"{int(float(payment))}|RUB",  
                                    'UF_CRM_1726492981056': self.data["дата рождения"], 
                                    'UF_CRM_1732785182099': self.start_date,
                                    #TEMP СДЕЛАТЬ ОБНОВЛЕНИЯ ОСТАЛЬНЫХ ПОЛЕЙ + реализовать многопоточность чтобы не лагало и не плакали сильно
                                }
                            }
                    response = requests.post(url, json=params, headers=HEADERS)
                    if response.status_code == 200:
                        self.log_signal.emit("Поля в битриксе заполнены")
                except Exception as e:
                    self.log_signal.emit(f"{e}")

            self.success_signal.emit(f"Документ успешно создан: {self.output_path}")
        except Exception as e:
            self.error_signal.emit(f"Ошибка: {e}")

class ContractGeneratorApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Генератор договора")
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        grid = QGridLayout()

        # Поля для данных
        
        grid.addWidget(QLabel("Номер договора\n(Только номер менеджера)"), 0, 0)
        self.dogovor_input = QLineEdit()
        grid.addWidget(self.dogovor_input, 0, 1)

        grid.addWidget(QLabel("ФИО:"), 1, 0)
        self.fio_input = QLineEdit()
        grid.addWidget(self.fio_input, 1, 1)

        grid.addWidget(QLabel("Дата рождения:"), 2, 0)
        self.dob_input = QDateEdit()
        self.dob_input.setDate(QDate.currentDate())
        grid.addWidget(self.dob_input, 2, 1)

        grid.addWidget(QLabel("Серия паспорта:"), 3, 0)
        self.seria_input = QLineEdit()
        grid.addWidget(self.seria_input, 3, 1)

        grid.addWidget(QLabel("Номер паспорта:"), 4, 0)
        self.nom_input = QLineEdit()
        grid.addWidget(self.nom_input, 4, 1)

        grid.addWidget(QLabel("Кем выдан:"), 5, 0)
        self.kem_input = QLineEdit()
        grid.addWidget(self.kem_input, 5, 1)

        grid.addWidget(QLabel("Дата выдачи:"), 6, 0)
        self.date_vidachi_input = QDateEdit()
        self.date_vidachi_input.setDate(QDate.currentDate())
        grid.addWidget(self.date_vidachi_input, 6, 1)

        grid.addWidget(QLabel("Код подразделения:"), 7, 0)
        self.kod_input = QLineEdit()
        grid.addWidget(self.kod_input, 7, 1)

        grid.addWidget(QLabel("Место рождения:"), 8, 0)
        self.mesto_rozhdeniya_input = QLineEdit()
        grid.addWidget(self.mesto_rozhdeniya_input, 8, 1)

        grid.addWidget(QLabel("Адрес регистрации:"), 9, 0)
        self.adres_input = QLineEdit()
        grid.addWidget(self.adres_input, 9, 1)

        grid.addWidget(QLabel("Телефон:"), 10, 0)
        self.telefon_input = QLineEdit()
        grid.addWidget(self.telefon_input, 10, 1)

        grid.addWidget(QLabel("Сумма (Юристы):"), 0, 2)
        self.summa_uristy_input = QLineEdit()
        grid.addWidget(self.summa_uristy_input, 0, 3)

        grid.addWidget(QLabel("Сумма (Бонус):"), 1, 2)
        self.summa_bonus_input = QLineEdit()
        grid.addWidget(self.summa_bonus_input, 1, 3)

        grid.addWidget(QLabel("Скидка:"), 2, 2)
        self.discount_input = QLineEdit()
        grid.addWidget(self.discount_input, 2, 3)

        grid.addWidget(QLabel("Дата рассрочки:"), 3, 2) #TEMP ПЕРЕДЕЛАТЬ РАССРОЧКУ А ТО КРИНГЕ И НЕ ЭЩКЕРЕ
        self.start_date_input = QDateEdit()
        self.start_date_input.setDate(QDate.currentDate())
        grid.addWidget(self.start_date_input, 3, 3)

        grid.addWidget(QLabel("Количество платежей:"), 4, 2)
        self.num_payments_input = QLineEdit()
        grid.addWidget(self.num_payments_input, 4, 3)
        
        grid.addWidget(QLabel("Сумма первого платежа:"), 5, 2)
        self.first_plateg_input = QLineEdit()
        grid.addWidget(self.first_plateg_input, 5, 3)
        


        layout.addLayout(grid)

        self.output_path_input = QLineEdit(self)
        self.output_path_input.setPlaceholderText("Выберите путь для сохранения...")
        self.output_path_button = QPushButton("Выбрать путь", self)
        self.output_path_button.clicked.connect(self.select_output_path)
        path_layout = QHBoxLayout()
        path_layout.addWidget(self.output_path_input)
        path_layout.addWidget(self.output_path_button)
        layout.addLayout(path_layout)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)

        self.generate_button = QPushButton("Создать документ")
        self.generate_button.clicked.connect(self.generate_contract)
        layout.addWidget(self.generate_button)

        self.setLayout(layout)

    def select_output_path(self):
        path, _ = QFileDialog.getSaveFileName(self, "Сохранить файл", "", "Документы Word (*.docx)")
        if path:
            self.output_path_input.setText(path)

    def log_message(self, message):
        self.log_text.append(message)

    def generate_contract(self):
        data = {
            "номер договора": self.dogovor_input.text(),
            "ФИО": self.fio_input.text(),
            "дата рождения": self.dob_input.date().toString("dd.MM.yyyy"),
            "серия": self.seria_input.text(),
            "номер": self.nom_input.text(),
            "кем": self.kem_input.text(),
            "дата выдачи": self.date_vidachi_input.date().toString("dd.MM.yyyy"),
            "код": self.kod_input.text(),
            "место рождения": self.mesto_rozhdeniya_input.text(),
            "адрес регистрации": self.adres_input.text(),
            "номер телефона": self.telefon_input.text(),
            "сумма юристы": self.summa_uristy_input.text(),
            "сумма бонус": self.summa_bonus_input.text(),
            "Первый платеж":self.first_plateg_input.text(),
            "data": datetime.today().strftime("%m/%Y"),
            "today": None,
            "words_sum":None,
        }

        missing_fields = [key for key, value in data.items() if not value and key != "today" and key != "words_sum"and key != "сумма бонус"]
        if missing_fields:
            self.log_message(f"Ошибка: Не заполнены поля: {', '.join(missing_fields)}")
            return

        try:
            template_path = resource_path('template.docx')
            output_path = self.output_path_input.text()
            discount = int(self.discount_input.text())
            num_payments = int(self.num_payments_input.text())
            start_date = self.start_date_input.date().toString("dd.MM.yyyy")

            self.worker = WorkerThread(data, template_path, output_path, discount, num_payments, start_date)
            self.worker.log_signal.connect(self.log_message)
            self.worker.success_signal.connect(lambda msg: QMessageBox.information(self, "Успех", msg))
            self.worker.error_signal.connect(lambda msg: QMessageBox.critical(self, "Ошибка", msg))
            self.worker.start()
            
        except Exception as e:
            self.log_message(f"Ошибка: {e}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ContractGeneratorApp()
    window.show()
    sys.exit(app.exec_())